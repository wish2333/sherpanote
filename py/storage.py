"""Local data persistence layer using SQLite with WAL mode.

Manages Record and Version CRUD, full-text search (FTS5),
multi-format export, and version history.
"""

from __future__ import annotations

import json
import sqlite3
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from py.config import _DEFAULT_DATA_DIR


class Storage:
    """SQLite-backed storage with atomic transactions and version tracking.

    All write operations are wrapped in transactions.
    Record objects use immutable design — save() returns new instances.
    """

    def __init__(self, db_path: str | None = None) -> None:
        if db_path is None:
            db_path = str(Path(_DEFAULT_DATA_DIR) / "data.db")
        Path(db_path).parent.mkdir(parents=True, exist_ok=True)
        self._db_path = db_path
        self._conn: sqlite3.Connection | None = None
        self._init_db()

    def _get_conn(self) -> sqlite3.Connection:
        if self._conn is None:
            self._conn = sqlite3.connect(self._db_path, check_same_thread=False)
            self._conn.execute("PRAGMA journal_mode=WAL")
            self._conn.execute("PRAGMA foreign_keys=ON")
            self._conn.row_factory = sqlite3.Row
        return self._conn

    def _init_db(self) -> None:
        """Create tables and indexes if they don't exist."""
        conn = self._get_conn()
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS records (
                id              TEXT PRIMARY KEY,
                title           TEXT NOT NULL DEFAULT '',
                audio_path      TEXT,
                transcript      TEXT NOT NULL DEFAULT '',
                segments_json   TEXT NOT NULL DEFAULT '[]',
                ai_results_json TEXT NOT NULL DEFAULT '{}',
                category        TEXT NOT NULL DEFAULT '',
                tags_json       TEXT NOT NULL DEFAULT '[]',
                duration_seconds REAL NOT NULL DEFAULT 0.0,
                created_at      TEXT NOT NULL,
                updated_at      TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS versions (
                record_id       TEXT NOT NULL,
                version         INTEGER NOT NULL,
                transcript      TEXT NOT NULL,
                segments_json   TEXT NOT NULL DEFAULT '[]',
                ai_results_json TEXT NOT NULL DEFAULT '{}',
                created_at      TEXT NOT NULL,
                PRIMARY KEY (record_id, version)
            );

            CREATE TABLE IF NOT EXISTS app_config (
                key   TEXT PRIMARY KEY,
                value TEXT NOT NULL
            );

            CREATE INDEX IF NOT EXISTS idx_records_created
                ON records(created_at DESC);

            CREATE INDEX IF NOT EXISTS idx_records_title
                ON records(title);

            -- Migration: add segments_json to versions table if missing
        """)
        # Migration: add segments_json column to versions table for older databases
        try:
            self._get_conn().execute(
                "ALTER TABLE versions ADD COLUMN segments_json TEXT NOT NULL DEFAULT '[]'"
            )
            self._get_conn().commit()
        except Exception:
            pass  # Column already exists
        conn.commit()

    # ---- Record CRUD ----

    def save(self, data: dict[str, Any]) -> dict[str, Any]:
        """Insert or update a record. Auto-generates id and timestamps.

        Does NOT create version snapshots -- use create_version() explicitly.
        For updates, provided fields are merged with existing record data
        to avoid unintended overwrites of unspecified fields.
        """
        conn = self._get_conn()
        now = datetime.now(timezone.utc).isoformat()

        record_id = data.get("id") or str(uuid.uuid4())
        is_update = bool(data.get("id"))

        if is_update:
            existing = self.get(record_id)
            if existing:
                # Merge: use existing values as defaults for any field not provided.
                merged = dict(existing)
                merged.update(data)
                data = merged

        segments_json = json.dumps(data.get("segments", []), ensure_ascii=False)
        ai_json = json.dumps(data.get("ai_results", {}), ensure_ascii=False)
        tags_json = json.dumps(data.get("tags", []), ensure_ascii=False)

        conn.execute(
            """INSERT INTO records (id, title, audio_path, transcript, segments_json,
               ai_results_json, category, tags_json, duration_seconds, created_at, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
               ON CONFLICT(id) DO UPDATE SET
                   title=excluded.title, audio_path=excluded.audio_path,
                   transcript=excluded.transcript, segments_json=excluded.segments_json,
                   ai_results_json=excluded.ai_results_json, category=excluded.category,
                   tags_json=excluded.tags_json, duration_seconds=excluded.duration_seconds,
                   updated_at=excluded.updated_at""",
            (
                record_id,
                data.get("title", ""),
                data.get("audio_path"),
                data.get("transcript", ""),
                segments_json,
                ai_json,
                data.get("category", ""),
                tags_json,
                data.get("duration_seconds", 0.0),
                data.get("created_at", now),
                now,
            ),
        )
        conn.commit()
        result = self.get(record_id)
        if result:
            result["version"] = self._get_current_version(record_id)
        return result  # type: ignore[return-value]

    def get(self, record_id: str) -> dict[str, Any] | None:
        """Fetch a single record by ID."""
        conn = self._get_conn()
        row = conn.execute("SELECT * FROM records WHERE id = ?", (record_id,)).fetchone()
        return self._row_to_dict(row) if row else None

    def list(self, filter: dict[str, Any] | None = None) -> list[dict[str, Any]]:
        """List records with optional filtering and sorting."""
        conn = self._get_conn()
        query = "SELECT * FROM records WHERE 1=1"
        params: list[Any] = []

        if filter:
            if filter.get("category"):
                query += " AND category = ?"
                params.append(filter["category"])
            if filter.get("keyword"):
                query += " AND (title LIKE ? OR transcript LIKE ?)"
                params.extend([f"%{filter['keyword']}%", f"%{filter['keyword']}%"])

        sort_by = (filter or {}).get("sort_by", "created_at")
        sort_order = (filter or {}).get("sort_order", "desc")
        if sort_by not in ("created_at", "updated_at", "title"):
            sort_by = "created_at"
        if sort_order not in ("asc", "desc"):
            sort_order = "desc"
        query += f" ORDER BY {sort_by} {sort_order}"

        rows = conn.execute(query, params).fetchall()
        return [self._row_to_dict(r) for r in rows]

    def delete(self, record_id: str) -> bool:
        """Delete a record and its versions."""
        conn = self._get_conn()
        conn.execute("DELETE FROM versions WHERE record_id = ?", (record_id,))
        conn.execute("DELETE FROM records WHERE id = ?", (record_id,))
        conn.commit()
        return conn.total_changes > 0

    # ---- Version History ----

    def get_versions(self, record_id: str) -> list[dict[str, Any]]:
        """Get all versions for a record, newest first."""
        conn = self._get_conn()
        rows = conn.execute(
            "SELECT * FROM versions WHERE record_id = ? ORDER BY version DESC",
            (record_id,),
        ).fetchall()
        return [self._row_to_dict(r) for r in rows]

    def _get_current_version(self, record_id: str) -> int:
        """Get the highest version number for a record."""
        conn = self._get_conn()
        row = conn.execute(
            "SELECT MAX(version) AS v FROM versions WHERE record_id = ?",
            (record_id,),
        ).fetchone()
        return row["v"] if row and row["v"] else 0

    def create_version(self, record_id: str, max_versions: int = 20) -> int:
        """Create a version snapshot of the current record state.

        Returns the new version number. If max_versions > 0, prunes
        the oldest versions to keep only the most recent max_versions.
        """
        conn = self._get_conn()
        record = self.get(record_id)
        if not record:
            raise ValueError(f"Record not found: {record_id}")

        now = datetime.now(timezone.utc).isoformat()
        current = self._get_current_version(record_id)
        new_version = current + 1

        conn.execute(
            "INSERT OR REPLACE INTO versions (record_id, version, transcript, segments_json, ai_results_json, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (record_id, new_version, record["transcript"], record.get("segments_json", "[]"), record["ai_results_json"], now),
        )
        conn.commit()

        if max_versions > 0:
            self._prune_versions(record_id, max_versions)

        return new_version

    def _prune_versions(self, record_id: str, max_versions: int) -> None:
        """Delete oldest versions, keeping only the most recent max_versions."""
        conn = self._get_conn()
        conn.execute(
            """DELETE FROM versions WHERE record_id = ? AND version NOT IN (
                SELECT version FROM versions WHERE record_id = ? ORDER BY version DESC LIMIT ?
            )""",
            (record_id, record_id, max_versions),
        )
        conn.commit()

    def delete_version(self, record_id: str, version: int) -> bool:
        """Delete a single version from a record's version history."""
        conn = self._get_conn()
        cursor = conn.execute(
            "DELETE FROM versions WHERE record_id = ? AND version = ?",
            (record_id, version),
        )
        conn.commit()
        return cursor.rowcount > 0

    def restore_version(self, record_id: str, version: int) -> dict[str, Any] | None:
        """Restore a record to a specific version (creates new version)."""
        conn = self._get_conn()
        row = conn.execute(
            "SELECT * FROM versions WHERE record_id = ? AND version = ?",
            (record_id, version),
        ).fetchone()
        if not row:
            return None

        old = self._row_to_dict(row)
        record = self.get(record_id)
        if not record:
            return None

        record["transcript"] = old["transcript"]
        record["segments_json"] = old.get("segments_json", "[]")
        record["segments"] = old.get("segments", [])
        record["ai_results_json"] = old["ai_results_json"]
        record["ai_results"] = old["ai_results"]
        return self.save(record)

    # ---- Export ----

    def export(self, record_id: str, fmt: str, output_dir: str | None = None, include_ai: bool = True) -> str:
        """Export a record to the specified format. Returns the output file path."""
        record = self.get(record_id)
        if not record:
            raise ValueError(f"Record not found: {record_id}")

        if output_dir is None:
            output_dir = str(Path(_DEFAULT_DATA_DIR) / "exports")

        Path(output_dir).mkdir(parents=True, exist_ok=True)
        title = record["title"] or "export"
        segments = json.loads(record.get("segments_json", "[]"))
        transcript = record.get("transcript", "")

        if fmt == "txt":
            path = self._export_txt(transcript, title, output_dir)
        elif fmt == "md":
            path = self._export_md(transcript, record, segments, title, output_dir, include_ai=include_ai)
        elif fmt == "srt":
            path = self._export_srt(segments, title, output_dir)
        elif fmt == "docx":
            path = self._export_docx(transcript, record, title, output_dir, include_ai=include_ai)
        else:
            raise ValueError(f"Unsupported format: {fmt}")

        return path

    @staticmethod
    def _export_txt(transcript: str, title: str, output_dir: str) -> str:
        path = str(Path(output_dir) / f"{title}.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"{title}\n{'=' * len(title)}\n\n{transcript}")
        return path

    @staticmethod
    def _export_md(transcript: str, record: dict, segments: list, title: str, output_dir: str, include_ai: bool = True) -> str:
        path = str(Path(output_dir) / f"{title}.md")
        lines = [f"# {title}\n"]
        if record.get("category"):
            lines.append(f"**Category:** {record['category']}")
        lines.append(f"**Created:** {record.get('created_at', '')}\n")
        lines.append("## Transcript\n")
        lines.append(transcript)
        if include_ai:
            ai_results = json.loads(record.get("ai_results_json", "{}"))
            for mode, content in ai_results.items():
                if content:
                    lines.append(f"\n## AI: {mode.title()}\n")
                    lines.append(content)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path

    @staticmethod
    def _export_srt(segments: list, title: str, output_dir: str) -> str:
        path = str(Path(output_dir) / f"{title}.srt")
        lines = []
        for i, seg in enumerate(segments, 1):
            start = Storage._seconds_to_srt_time(seg.get("start_time", 0))
            end = Storage._seconds_to_srt_time(seg.get("end_time", 0))
            lines.append(f"{i}")
            lines.append(f"{start} --> {end}")
            lines.append(seg.get("text", ""))
            lines.append("")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return path

    @staticmethod
    def _export_docx(transcript: str, record: dict, title: str, output_dir: str, include_ai: bool = True) -> str:
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=1)
        if record.get("category"):
            doc.add_paragraph(f"Category: {record['category']}")
        doc.add_paragraph(f"Created: {record.get('created_at', '')}")
        doc.add_heading("Transcript", level=2)
        doc.add_paragraph(transcript)

        if include_ai:
            ai_results = json.loads(record.get("ai_results_json", "{}"))
            for mode, content in ai_results.items():
                if content:
                    doc.add_heading(f"AI: {mode.title()}", level=2)
                    doc.add_paragraph(content)

        path = str(Path(output_dir) / f"{title}.docx")
        doc.save(path)
        return path

    @staticmethod
    def _seconds_to_srt_time(seconds: float) -> str:
        """Convert seconds to SRT time format HH:MM:SS,mmm."""
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        ms = int((seconds % 1) * 1000)
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

    # ---- Helpers ----

    @staticmethod
    def _row_to_dict(row: sqlite3.Row) -> dict[str, Any]:
        """Convert a sqlite3.Row to a plain dict, parsing JSON fields."""
        d = dict(row)
        json_defaults = {
            "segments_json": [],
            "ai_results_json": {},
            "tags_json": [],
        }
        for json_field, default in json_defaults.items():
            if json_field in d:
                try:
                    d[json_field.replace("_json", "")] = json.loads(d[json_field])
                except (json.JSONDecodeError, TypeError):
                    d[json_field.replace("_json", "")] = default
        return d

    def close(self) -> None:
        if self._conn:
            try:
                self._conn.execute("PRAGMA wal_checkpoint(TRUNCATE)")
            except Exception:
                pass
            self._conn.close()
            self._conn = None
